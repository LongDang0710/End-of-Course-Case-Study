import tkinter as tk
from tkinter import ttk, colorchooser, font, messagebox, filedialog
import sqlite3
import csv
import pandas as pd
from docx import Document
from fpdf import FPDF


class LibraryManagementSystem:
    def __init__(self, root):
        self.root = root
        self.root.title("Library Management System")
        self.root.geometry("900x700")

        # Kết nối tới cơ sở dữ liệu SQLite
        self.conn = sqlite3.connect('library.db')
        self.cursor = self.conn.cursor()
        self.create_tables()

        self.create_toolbar()
        self.create_notebook()
        self.create_tabs()

    def create_toolbar(self):
        # Tạo thanh công cụ và thêm nút File với menu con
        toolbar = tk.Frame(self.root, bd=1, relief=tk.RAISED)
        toolbar.pack(side=tk.TOP, fill=tk.X)

        # Nút File với menu con (gồm các nút Save As, Reset All, và Import)
        file_button = tk.Menubutton(toolbar, text="File", relief=tk.RAISED)
        file_button.pack(side=tk.LEFT, padx=2, pady=2)
        file_menu = tk.Menu(file_button, tearoff=0)
        file_menu.add_command(label="Save As", command=self.open_save_as_window)
        file_menu.add_command(label="Reset All", command=self.reset_all)
        file_menu.add_command(label="Import", command=self.import_data)
        file_button.config(menu=file_menu)

    def create_notebook(self):
        # Tạo notebook (giao diện tab)
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(pady=10, expand=True)

    def create_tabs(self):
        # Setup các tabs
        self.books_tab = ttk.Frame(self.notebook, width=800, height=600)
        self.members_tab = ttk.Frame(self.notebook, width=800, height=600)
        self.transactions_tab = ttk.Frame(self.notebook, width=800, height=600)
        self.settings_tab = ttk.Frame(self.notebook, width=800, height=600)

        self.notebook.add(self.books_tab, text="Books")
        self.notebook.add(self.members_tab, text="Members")
        self.notebook.add(self.transactions_tab, text="Transactions")
        self.notebook.add(self.settings_tab, text="Settings")

        self.setup_books_tab()
        self.setup_members_tab()
        self.setup_transactions_tab()
        self.setup_settings_tab()

    def setup_books_tab(self):
        # Setup thông tin tab Books
        self.books_tree = self.create_tree_view(self.books_tab, ["Title", "Author", "Genre", "Quantity", "Available"])
        self.create_book_form(self.books_tab)
        self.load_books()
        tk.Button(self.books_tab, text="Reset Books", command=self.reset_books).pack(pady=10)

    def setup_members_tab(self):
        # Setup thông tin tab Members
        self.members_tree = self.create_tree_view(self.members_tab,
                                                  ["Member ID", "Name", "Membership Date", "Books Borrowed",
                                                   "Quantity Borrowed"])
        self.create_member_form(self.members_tab)
        self.load_members()
        tk.Button(self.members_tab, text="Reset Members", command=self.reset_members).pack(pady=10)

    def setup_transactions_tab(self):
        # Setup thông tin tab Transactions
        self.transactions_tree = self.create_tree_view(self.transactions_tab,
                                                       ["Transaction ID", "Book ID", "Member ID", "Borrow Date",
                                                        "Return Date"])
        self.create_transaction_form(self.transactions_tab)
        self.load_transactions()
        tk.Button(self.transactions_tab, text="Reset Transactions", command=self.reset_transactions).pack(pady=10)

    def setup_settings_tab(self):
        # Setup thông tin tab Settings
        self.create_settings_form(self.settings_tab)

    # Tạo tính năng Reset All
    def reset_all(self):
        if messagebox.askyesno("Warning",
                               "This action will reset everything to default and cannot be undone. Do you want to proceed?"):
            self.reset_books(confirm=False)
            self.reset_members(confirm=False)
            self.reset_transactions(confirm=False)
            self.reset_settings(confirm=False)
            messagebox.showinfo("Reset All", "Reset Successfully!")

    # Tạo tính năng reset dữ liệu sách
    def reset_books(self, confirm=True):
        if not confirm or messagebox.askyesno("Warning",
                                              "This action will delete all book's data. Do you want to proceed?"):
            self.cursor.execute('DELETE FROM books')
            self.conn.commit()
            for item in self.books_tree.get_children():
                self.books_tree.delete(item)
            if confirm:
                messagebox.showinfo("Books Data Reset", "Reset successfully!")

    # Tạo tính năng reset dữ liệu thành viên
    def reset_members(self, confirm=True):
        if not confirm or messagebox.askyesno("Warning",
                                              "This action will delete all member's data. Do you want to proceed?"):
            self.cursor.execute('DELETE FROM members')
            self.conn.commit()
            for item in self.members_tree.get_children():
                self.members_tree.delete(item)
            if confirm:
                messagebox.showinfo("Members Data Reset", "Reset successfully!")

    # Tạo tính năng reset dữ liệu giao dịch
    def reset_transactions(self, confirm=True):
        if not confirm or messagebox.askyesno("Warning",
                                              "This action will delete all transaction's data. Do you want to proceed?"):
            self.cursor.execute('DELETE FROM transactions')
            self.conn.commit()
            for item in self.transactions_tree.get_children():
                self.transactions_tree.delete(item)
            if confirm:
                messagebox.showinfo("Transactions Data Reset", "Reset successfully!")

    # Tạo tính năng reset cài đặt
    def reset_settings(self, confirm=True):
        if not confirm or messagebox.askyesno("Warning",
                                              "This action will reset all settings to default. Do you want to proceed?"):
            self.bg_color_var.set('#FFFFFF')
            self.font_size_var.set(10)
            self.font_style_var.set('TkDefaultFont')
            print(
                f"Reset settings to default: bg_color={self.bg_color_var.get()}, font_size={self.font_size_var.get()}, font_style={self.font_style_var.get()}")
            self.apply_settings()
            if confirm:
                messagebox.showinfo("Settings Reset", "Settings reset successfully!")

    # Tạo bảng trong cơ sở dữ liệu
    def create_tables(self):
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS books (
                title TEXT PRIMARY KEY,
                author TEXT,
                genre TEXT,
                quantity INTEGER,
                available INTEGER
            )
        ''')
        self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS members (
                    member_id TEXT PRIMARY KEY,
                    name TEXT,
                    membership_date TEXT,
                    books_borrowed TEXT,
                    quantity_borrowed INTEGER
                )
            ''')
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS transactions (
                transaction_id TEXT PRIMARY KEY,
                book_id TEXT,
                member_id TEXT,
                borrow_date TEXT,
                return_date TEXT
            )
        ''')
        self.conn.commit()

    # Tạo TreeView để hiện thị dữ liệu
    def create_tree_view(self, parent, columns):
        tree_frame = tk.Frame(parent)
        tree_frame.pack(pady=20)
        tree_scroll = tk.Scrollbar(tree_frame)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        tree = ttk.Treeview(tree_frame, yscrollcommand=tree_scroll.set, columns=columns, show="headings")
        tree.pack()
        tree_scroll.config(command=tree.yview)
        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=150)
        return tree

    # Thiết kết tab Books
    def create_book_form(self, parent):
        form_frame = tk.Frame(parent)
        form_frame.pack(pady=20)
        tk.Label(form_frame, text="Title").grid(row=0, column=0, padx=10, pady=5)
        self.title_entry = tk.Entry(form_frame)
        self.title_entry.grid(row=0, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Author").grid(row=1, column=0, padx=10, pady=5)
        self.author_entry = tk.Entry(form_frame)
        self.author_entry.grid(row=1, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Genre").grid(row=2, column=0, padx=10, pady=5)
        self.genre_entry = tk.Entry(form_frame)
        self.genre_entry.grid(row=2, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Quantity").grid(row=3, column=0, padx=10, pady=5)
        self.quantity_entry = tk.Entry(form_frame)
        self.quantity_entry.grid(row=3, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Available").grid(row=4, column=0, padx=10, pady=5)
        self.available_entry = tk.Entry(form_frame)
        self.available_entry.grid(row=4, column=1, padx=10, pady=5)
        tk.Button(form_frame, text="Add Book", command=self.add_book).grid(row=5, column=0, padx=10, pady=10)
        tk.Button(form_frame, text="Update Book", command=self.update_book).grid(row=5, column=1, padx=10, pady=10)
        tk.Button(form_frame, text="Delete Book", command=self.delete_book).grid(row=6, column=0, padx=10, pady=10)
        tk.Button(form_frame, text="Clear Fields", command=self.clear_fields).grid(row=6, column=1, padx=10, pady=10)

    # Thiết kết tab Members
    def create_member_form(self, parent):
        form_frame = tk.Frame(parent)
        form_frame.pack(pady=20)
        tk.Label(form_frame, text="Member ID").grid(row=0, column=0, padx=10, pady=5)
        self.member_id_entry = tk.Entry(form_frame)
        self.member_id_entry.grid(row=0, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Name").grid(row=1, column=0, padx=10, pady=5)
        self.name_entry = tk.Entry(form_frame)
        self.name_entry.grid(row=1, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Membership Date").grid(row=2, column=0, padx=10, pady=5)
        self.membership_date_entry = tk.Entry(form_frame)
        self.membership_date_entry.grid(row=2, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Books Borrowed").grid(row=3, column=0, padx=10, pady=5)
        self.books_borrowed_entry = tk.Entry(form_frame)
        self.books_borrowed_entry.grid(row=3, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Quantity Borrowed").grid(row=4, column=0, padx=10, pady=5)
        self.quantity_borrowed_entry = tk.Entry(form_frame)
        self.quantity_borrowed_entry.grid(row=4, column=1, padx=10, pady=5)
        tk.Button(form_frame, text="Add Member", command=self.add_member).grid(row=5, column=0, padx=10, pady=10)
        tk.Button(form_frame, text="Update Member", command=self.update_member).grid(row=5, column=1, padx=10, pady=10)
        tk.Button(form_frame, text="Delete Member", command=self.delete_member).grid(row=6, column=0, padx=10, pady=10)
        tk.Button(form_frame, text="Clear Fields", command=self.clear_member_fields).grid(row=6, column=1, padx=10,
                                                                                          pady=10)

    # Thiết kế tab Transactions
    def create_transaction_form(self, parent):
        form_frame = tk.Frame(parent)
        form_frame.pack(pady=20)
        tk.Label(form_frame, text="Transaction ID").grid(row=0, column=0, padx=10, pady=5)
        self.transaction_id_entry = tk.Entry(form_frame)
        self.transaction_id_entry.grid(row=0, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Book ID").grid(row=1, column=0, padx=10, pady=5)
        self.trans_book_id_entry = tk.Entry(form_frame)
        self.trans_book_id_entry.grid(row=1, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Member ID").grid(row=2, column=0, padx=10, pady=5)
        self.trans_member_id_entry = tk.Entry(form_frame)
        self.trans_member_id_entry.grid(row=2, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Borrow Date").grid(row=3, column=0, padx=10, pady=5)
        self.borrow_date_entry = tk.Entry(form_frame)
        self.borrow_date_entry.grid(row=3, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Return Date").grid(row=4, column=0, padx=10, pady=5)
        self.return_date_entry = tk.Entry(form_frame)
        self.return_date_entry.grid(row=4, column=1, padx=10, pady=5)
        tk.Button(form_frame, text="Add Transaction", command=self.add_transaction).grid(row=5, column=0, padx=10,
                                                                                         pady=10)
        tk.Button(form_frame, text="Update Transaction", command=self.update_transaction).grid(row=5, column=1, padx=10,
                                                                                               pady=10)
        tk.Button(form_frame, text="Delete Transaction", command=self.delete_transaction).grid(row=6, column=0, padx=10,
                                                                                               pady=10)
        tk.Button(form_frame, text="Clear Fields", command=self.clear_transaction_fields).grid(row=6, column=1, padx=10,
                                                                                               pady=10)

    # Thiết kế tab Settings
    def create_settings_form(self, parent):
        form_frame = tk.Frame(parent)
        form_frame.pack(pady=20)
        tk.Label(form_frame, text="Background Color").grid(row=0, column=0, padx=10, pady=5)
        self.bg_color_var = tk.StringVar()
        self.bg_color_button = tk.Button(form_frame, text="Choose Color", command=self.choose_bg_color)
        self.bg_color_button.grid(row=0, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Font Size").grid(row=1, column=0, padx=10, pady=5)
        self.font_size_var = tk.IntVar(value=10)
        self.font_size_spinbox = tk.Spinbox(form_frame, from_=8, to=20, textvariable=self.font_size_var)
        self.font_size_spinbox.grid(row=1, column=1, padx=10, pady=5)
        tk.Label(form_frame, text="Font Style").grid(row=2, column=0, padx=10, pady=5)
        self.font_style_var = tk.StringVar()
        self.font_style_combobox = ttk.Combobox(form_frame, textvariable=self.font_style_var, values=font.families())
        self.font_style_combobox.grid(row=2, column=1, padx=10, pady=5)
        tk.Button(form_frame, text="Apply", command=self.apply_settings).grid(row=3, column=0, columnspan=2, pady=10)
        tk.Button(form_frame, text="Reset Settings", command=self.reset_settings).grid(row=4, column=0, columnspan=2,
                                                                                       pady=10)

    # Thiết kế các tính năng cho hệ thống
    # Màu nền và phông chữ
    def choose_bg_color(self):
        color = colorchooser.askcolor()[1]
        if color:
            self.bg_color_var.set(color)

    # Áp dụng các cài đặt
    def apply_settings(self):
        # Màu nền
        bg_color = self.bg_color_var.get()
        if bg_color:
            self.apply_background_color(self.root, bg_color)

        # Cỡ chữ và phông chữ
        font_size = self.font_size_var.get()
        font_style = self.font_style_var.get()
        if font_style:
            custom_font = font.Font(family=font_style, size=font_size)
            self.apply_font(self.root, custom_font)

    # Mô tả function áp dụng cài đặt
    # Màu nền
    def apply_background_color(self, widget, color):
        # Liệt kê các widgets có hỗ trợ màu nền
        supported_widgets = (tk.Frame, tk.Label, tk.Button, tk.Entry, ttk.Treeview, tk.Text)
        if isinstance(widget, supported_widgets):
            try:
                widget.configure(bg=color)
            except tk.TclError:
                pass
        for child in widget.winfo_children():
            self.apply_background_color(child, color)

    # Phông chữ
    def apply_font(self, widget, custom_font):
        try:
            widget.configure(font=custom_font)
        except tk.TclError:
            pass
        for child in widget.winfo_children():
            self.apply_font(child, custom_font)

    # Mô tả các function trong các tabs đã tạo
    # Thêm sách
    def add_book(self):
        title = self.title_entry.get()
        author = self.author_entry.get()
        genre = self.genre_entry.get()
        quantity = self.quantity_entry.get()
        available = self.available_entry.get()
        if title and author and genre and quantity and available:
            try:
                self.cursor.execute(
                    'INSERT INTO books (title, author, genre, quantity, available) VALUES (?, ?, ?, ?, ?)',
                    (title, author, genre, quantity, available))
                self.conn.commit()
                self.books_tree.insert("", "end", values=(title, author, genre, quantity, available))
                self.clear_fields()
            except sqlite3.IntegrityError:
                messagebox.showwarning("Warning", "Book with this title already exists!")
        else:
            messagebox.showwarning("Warning", "All fields are required, please try again!")

    # Cập nhật sách
    def update_book(self):
        selected_item = self.books_tree.selection()
        if selected_item:
            title = self.title_entry.get()
            author = self.author_entry.get()
            genre = self.genre_entry.get()
            quantity = self.quantity_entry.get()
            available = self.available_entry.get()
            self.books_tree.item(selected_item, values=(title, author, genre, quantity, available))
            self.cursor.execute('UPDATE books SET author = ?, genre = ?, quantity = ?, available = ? WHERE title = ?',
                                (author, genre, quantity, available, title))
            self.conn.commit()
            self.clear_fields()
        else:
            messagebox.showwarning("Warning", "You must select a book!")

    # Xoá sách
    def delete_book(self):
        selected_item = self.books_tree.selection()
        if selected_item:
            title = self.books_tree.item(selected_item, 'values')[0]
            self.books_tree.delete(selected_item)
            self.cursor.execute('DELETE FROM books WHERE title = ?', (title,))
            self.conn.commit()
            self.clear_fields()
        else:
            messagebox.showwarning("Warning", "You must select a book!")

    # Xoá tất cả các trường thông tin sách vừa được nhập vào
    def clear_fields(self):
        self.title_entry.delete(0, tk.END)
        self.author_entry.delete(0, tk.END)
        self.genre_entry.delete(0, tk.END)
        self.quantity_entry.delete(0, tk.END)
        self.available_entry.delete(0, tk.END)

    # Thêm thành viên
    def add_member(self):
        member_id = self.member_id_entry.get()
        name = self.name_entry.get()
        membership_date = self.membership_date_entry.get()
        books_borrowed = self.books_borrowed_entry.get()
        quantity_borrowed = self.quantity_borrowed_entry.get()
        if member_id and name and membership_date and books_borrowed and quantity_borrowed:
            try:
                self.cursor.execute('INSERT INTO members VALUES (?, ?, ?, ?, ?)',
                                    (member_id, name, membership_date, books_borrowed, quantity_borrowed))
                self.conn.commit()
                self.members_tree.insert("", "end",
                                         values=(member_id, name, membership_date, books_borrowed, quantity_borrowed))
                self.clear_member_fields()
            except sqlite3.IntegrityError:
                messagebox.showwarning("Warning", "Member with this ID already exists!")
        else:
            messagebox.showwarning("Warning", "All fields are required, please try again!")

    # Cập nhật thành viên
    def update_member(self):
        selected_item = self.members_tree.selection()
        if selected_item:
            member_id = self.member_id_entry.get()
            name = self.name_entry.get()
            membership_date = self.membership_date_entry.get()
            books_borrowed = self.books_borrowed_entry.get()
            quantity_borrowed = self.quantity_borrowed_entry.get()
            self.members_tree.item(selected_item,
                                   values=(member_id, name, membership_date, books_borrowed, quantity_borrowed))
            self.cursor.execute(
                'UPDATE members SET name = ?, membership_date = ?, books_borrowed = ?, quantity_borrowed = ? WHERE member_id = ?',
                (name, membership_date, books_borrowed, quantity_borrowed, member_id))
            self.conn.commit()
            self.clear_member_fields()
        else:
            messagebox.showwarning("Warning", "You must select a member!")

    # Xoá thành viên
    def delete_member(self):
        selected_item = self.members_tree.selection()
        if selected_item:
            member_id = self.members_tree.item(selected_item, 'values')[0]
            self.members_tree.delete(selected_item)
            self.cursor.execute('DELETE FROM members WHERE member_id = ?', (member_id,))
            self.conn.commit()
            self.clear_member_fields()
        else:
            messagebox.showwarning("Warning", "You must select a member!")

    # Xoá tất cả các trường thông tin thành viên vừa được nhập vào
    def clear_member_fields(self):
        self.member_id_entry.delete(0, tk.END)
        self.name_entry.delete(0, tk.END)
        self.membership_date_entry.delete(0, tk.END)
        self.books_borrowed_entry.delete(0, tk.END)

    # Thêm giao dịch
    def add_transaction(self):
        transaction_id = self.transaction_id_entry.get()
        book_id = self.trans_book_id_entry.get()
        member_id = self.trans_member_id_entry.get()
        borrow_date = self.borrow_date_entry.get()
        return_date = self.return_date_entry.get()
        if transaction_id and book_id and member_id and borrow_date and return_date:
            try:
                self.cursor.execute('INSERT INTO transactions VALUES (?, ?, ?, ?, ?)',
                                    (transaction_id, book_id, member_id, borrow_date, return_date))
                self.conn.commit()
                self.transactions_tree.insert("", "end",
                                              values=(transaction_id, book_id, member_id, borrow_date, return_date))
                self.clear_transaction_fields()
            except sqlite3.IntegrityError:
                messagebox.showwarning("Warning", "Transaction with this ID already exists!")
        else:
            messagebox.showwarning("Warning", "All fields are required, please try again!")

    # Cập nhật giao dịch
    def update_transaction(self):
        selected_item = self.transactions_tree.selection()
        if selected_item:
            transaction_id = self.transaction_id_entry.get()
            book_id = self.trans_book_id_entry.get()
            member_id = self.trans_member_id_entry.get()
            borrow_date = self.borrow_date_entry.get()
            return_date = self.return_date_entry.get()
            self.transactions_tree.item(selected_item,
                                        values=(transaction_id, book_id, member_id, borrow_date, return_date))
            self.cursor.execute(
                'UPDATE transactions SET book_id = ?, member_id = ?, borrow_date = ?, return_date = ? WHERE transaction_id = ?',
                (book_id, member_id, borrow_date, return_date, transaction_id))
            self.conn.commit()
            self.clear_transaction_fields()
        else:
            messagebox.showwarning("Warning", "You must select a transaction!")

    # Xoá giao dịch
    def delete_transaction(self):
        selected_item = self.transactions_tree.selection()
        if selected_item:
            transaction_id = self.transactions_tree.item(selected_item, 'values')[0]
            self.transactions_tree.delete(selected_item)
            self.cursor.execute('DELETE FROM transactions WHERE transaction_id = ?', (transaction_id,))
            self.conn.commit()
            self.clear_transaction_fields()
        else:
            messagebox.showwarning("Warning", "You must select a transaction!")

    # Xoá tất cả các trường thông tin giao dịch vừa được nhập vào
    def clear_transaction_fields(self):
        self.transaction_id_entry.delete(0, tk.END)
        self.trans_book_id_entry.delete(0, tk.END)
        self.trans_member_id_entry.delete(0, tk.END)
        self.borrow_date_entry.delete(0, tk.END)
        self.return_date_entry.delete(0, tk.END)

    # Xem sách
    def load_books(self):
        self.cursor.execute('SELECT title, author, genre, quantity, available FROM books')
        rows = self.cursor.fetchall()
        for row in rows:
            self.books_tree.insert("", "end", values=row)

    # Xem thành viên
    def load_members(self):
        self.cursor.execute('SELECT member_id, name, membership_date, books_borrowed, quantity_borrowed FROM members')
        rows = self.cursor.fetchall()
        for row in rows:
            self.members_tree.insert("", "end", values=row)

    # Xem giao dịch
    def load_transactions(self):
        self.cursor.execute('SELECT * FROM transactions')
        rows = self.cursor.fetchall()
        for row in rows:
            self.transactions_tree.insert("", "end", values=row)

    # Thiết kế nút và cửa sổ Save as
    def open_save_as_window(self):
        self.save_as_window = tk.Toplevel(self.root)
        self.save_as_window.title("Save As")
        self.save_as_window.geometry("400x300")
        tk.Label(self.save_as_window, text="Save Format").pack(pady=10)
        self.save_format_var = tk.StringVar(value="sqlite")
        formats = ["SQLite", "Excel", "Word", "PDF"]
        for fmt in formats:
            tk.Radiobutton(self.save_as_window, text=fmt, variable=self.save_format_var, value=fmt.lower()).pack(
                anchor=tk.W)
        tk.Label(self.save_as_window, text="Select Data to Save").pack(pady=10)
        self.save_books_var = tk.BooleanVar()
        self.save_members_var = tk.BooleanVar()
        self.save_transactions_var = tk.BooleanVar()
        tk.Checkbutton(self.save_as_window, text="Books", variable=self.save_books_var).pack(anchor=tk.W)
        tk.Checkbutton(self.save_as_window, text="Members", variable=self.save_members_var).pack(anchor=tk.W)
        tk.Checkbutton(self.save_as_window, text="Transactions", variable=self.save_transactions_var).pack(anchor=tk.W)
        tk.Button(self.save_as_window, text="Save", command=self.save_as).pack(pady=10)
        tk.Button(self.save_as_window, text="Cancel", command=self.save_as_window.destroy).pack(pady=10)

    # Mô tả tính năng Save as
    def save_as(self):
        format_selected = self.save_format_var.get()
        if format_selected not in ["sqlite", "excel", "word", "pdf"]:
            messagebox.showwarning("Warning", "Unknown save format selected")
            return
        if not any([self.save_books_var.get(), self.save_members_var.get(), self.save_transactions_var.get()]):
            messagebox.showwarning("Warning", "No data selected to save")
            return

        # Chuẩn hoá định dạng lưu
        extensions = {"sqlite": "db", "excel": "xlsx", "word": "docx", "pdf": "pdf"}
        file_path = filedialog.asksaveasfilename(defaultextension=f".{extensions[format_selected]}",
                                                 filetypes=[(format_selected.capitalize(),
                                                             f"*.{extensions[format_selected]}")])
        if not file_path:
            return

        save_successful = True
        if self.save_books_var.get():
            save_successful &= self.save_books(format_selected, file_path)
        if self.save_members_var.get():
            save_successful &= self.save_members(format_selected, file_path)
        if self.save_transactions_var.get():
            save_successful &= self.save_transactions(format_selected, file_path)

        if save_successful:
            messagebox.showinfo("Save As", "Data saved successfully")
        else:
            messagebox.showerror("Error", "Failed to save some data")
        self.save_as_window.destroy()

    # Mô tả các định dạng và dữ liệu lưu (sqlite, excel, word, pdf)
    # Sách
    def save_books(self, format_selected, file_path):
        try:
            if format_selected == "sqlite":
                self.save_books_as_sqlite(file_path)
            elif format_selected == "excel":
                self.save_books_as_excel(file_path)
            elif format_selected == "word":
                self.save_books_as_word(file_path)
            elif format_selected == "pdf":
                self.save_books_as_pdf(file_path)
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save books: {str(e)}")
            return False

    # Thành viên
    def save_members(self, format_selected, file_path):
        try:
            if format_selected == "sqlite":
                self.save_members_as_sqlite(file_path)
            elif format_selected == "excel":
                self.save_members_as_excel(file_path)
            elif format_selected == "word":
                self.save_members_as_word(file_path)
            elif format_selected == "pdf":
                self.save_members_as_pdf(file_path)
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save members: {str(e)}")
            return False

    # Giao dịch
    def save_transactions(self, format_selected, file_path):
        try:
            if format_selected == "sqlite":
                self.save_transactions_as_sqlite(file_path)
            elif format_selected == "excel":
                self.save_transactions_as_excel(file_path)
            elif format_selected == "word":
                self.save_transactions_as_word(file_path)
            elif format_selected == "pdf":
                self.save_transactions_as_pdf(file_path)
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save transactions: {str(e)}")
            return False

    # Mô tả tính năng lưu cho các dữ liệu trong từng tab
    # Sách
    def save_books_as_sqlite(self, file_path):
        # Lưu dạng Sqlite
        pass

    def save_books_as_excel(self, file_path):
        import pandas as pd
        data = []
        self.cursor.execute('SELECT title, author, genre, quantity, available FROM books')
        rows = self.cursor.fetchall()
        for row in rows:
            data.append(row)
        df = pd.DataFrame(data, columns=["Title", "Author", "Genre", "Quantity", "Available"])
        df.to_excel(file_path, index=False)

    def save_books_as_word(self, file_path):
        from docx import Document
        self.cursor.execute('SELECT title, author, genre, quantity, available FROM books')
        rows = self.cursor.fetchall()
        doc = Document()
        doc.add_heading('Books Data', 0)
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Title'
        hdr_cells[1].text = 'Author'
        hdr_cells[2].text = 'Genre'
        hdr_cells[3].text = 'Quantity'
        hdr_cells[4].text = 'Available'
        for row in rows:
            cells = table.add_row().cells
            for i, item in enumerate(row):
                cells[i].text = str(item)
        doc.save(file_path)

    def save_books_as_pdf(self, file_path):
        from fpdf import FPDF
        self.cursor.execute('SELECT title, author, genre, quantity, available FROM books')
        rows = self.cursor.fetchall()
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Books Data", ln=True, align='C')
        for row in rows:
            line = ", ".join(map(str, row))
            pdf.cell(200, 10, txt=line, ln=True)
        pdf.output(file_path)

    # Thành viên
    def save_members_as_sqlite(self, file_path):
        pass

    def save_members_as_excel(self, file_path):
        data = []
        self.cursor.execute('SELECT member_id, name, membership_date, books_borrowed, quantity_borrowed FROM members')
        rows = self.cursor.fetchall()
        for row in rows:
            data.append(row)
        df = pd.DataFrame(data, columns=["Member ID", "Name", "Membership Date", "Books Borrowed", "Quantity Borrowed"])
        df.to_excel(file_path, index=False)

    def save_members_as_word(self, file_path):
        data = []
        self.cursor.execute('SELECT member_id, name, membership_date, books_borrowed, quantity_borrowed FROM members')
        rows = self.cursor.fetchall()
        doc = Document()
        doc.add_heading('Members Data', 0)
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Member ID'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Membership Date'
        hdr_cells[3].text = 'Books Borrowed'
        hdr_cells[4].text = 'Quantity Borrowed'
        for row in rows:
            cells = table.add_row().cells
            for i, item in enumerate(row):
                cells[i].text = str(item)
        doc.save(file_path)

    def save_members_as_pdf(self, file_path):
        data = []
        self.cursor.execute('SELECT member_id, name, membership_date, books_borrowed, quantity_borrowed FROM members')
        rows = self.cursor.fetchall()
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Members Data", ln=True, align='C')
        for row in rows:
            line = ", ".join(map(str, row))
            pdf.cell(200, 10, txt=line, ln=True)
        pdf.output(file_path)

    # Giao dịch
    def save_transactions_as_sqlite(self, file_path):
        pass

    def save_transactions_as_excel(self, file_path):
        data = []
        self.cursor.execute('SELECT transaction_id, book_id, member_id, borrow_date, return_date FROM transactions')
        rows = self.cursor.fetchall()
        for row in rows:
            data.append(row)
        df = pd.DataFrame(data, columns=["Transaction ID", "Book ID", "Member ID", "Borrow Date", "Return Date"])
        df.to_excel(file_path, index=False)

    def save_transactions_as_word(self, file_path):
        data = []
        self.cursor.execute('SELECT transaction_id, book_id, member_id, borrow_date, return_date FROM transactions')
        rows = self.cursor.fetchall()
        doc = Document()
        doc.add_heading('Transactions Data', 0)
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Transaction ID'
        hdr_cells[1].text = 'Book ID'
        hdr_cells[2].text = 'Member ID'
        hdr_cells[3].text = 'Borrow Date'
        hdr_cells[4].text = 'Return Date'
        for row in rows:
            cells = table.add_row().cells
            for i, item in enumerate(row):
                cells[i].text = str(item)
        doc.save(file_path)

    def save_transactions_as_pdf(self, file_path):
        data = []
        self.cursor.execute('SELECT transaction_id, book_id, member_id, borrow_date, return_date FROM transactions')
        rows = self.cursor.fetchall()
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Transactions Data", ln=True, align='C')
        for row in rows:
            line = ", ".join(map(str, row))
            pdf.cell(200, 10, txt=line, ln=True)
        pdf.output(file_path)

    # Mô tả tính năng Import (Nhập dữ liệu từ file csv)
    def import_data(self):
        file_path = filedialog.askopenfilename(defaultextension=".csv", filetypes=[("CSV files", "*.csv")])
        if not file_path:
            return
        with open(file_path, newline='') as csvfile:
            reader = csv.reader(csvfile)
            data = list(reader)
        if not data:
            messagebox.showwarning("Warning", "No data found in the selected file")
            return
        headers = data[0]
        if headers == ["Title", "Author", "Genre", "Quantity", "Available"]:
            self.import_books(data[1:])
        elif headers == ["Member ID", "Name", "Membership Date", "Books Borrowed", "Quantity Borrowed"]:
            self.import_members(data[1:])
        elif headers == ["Transaction ID", "Book ID", "Member ID", "Borrow Date", "Return Date"]:
            self.import_transactions(data[1:])
        else:
            messagebox.showwarning("Warning", "Invalid data format in the selected file")

    # Nhập dữ liệu sách
    def import_books(self, data):
        self.reset_books(confirm=False)
        for row in data:
            self.cursor.execute('INSERT INTO books (title, author, genre, quantity, available) VALUES (?, ?, ?, ?, ?)',
                                row)
        self.conn.commit()
        self.load_books()
        messagebox.showinfo("Import Data", "Books data imported successfully")

    # Nhập dữ liệu thành viên
    def import_members(self, data):
        self.reset_members(confirm=False)
        for row in data:
            self.cursor.execute(
                'INSERT INTO members (member_id, name, membership_date, books_borrowed, quantity_borrowed) VALUES (?, ?, ?, ?, ?)',
                row)
        self.conn.commit()
        self.load_members()
        messagebox.showinfo("Import Data", "Members data imported successfully")

    # Nhập dữ liệu giao dịch
    def import_transactions(self, data):
        self.reset_transactions(confirm=False)
        for row in data:
            self.cursor.execute(
                'INSERT INTO transactions (transaction_id, book_id, member_id, borrow_date, return_date) VALUES (?, ?, ?, ?, ?)',
                row)
        self.conn.commit()
        self.load_transactions()
        messagebox.showinfo("Import Data", "Transactions data imported successfully")


# Chạy ứng dụng
if __name__ == "__main__":
    root = tk.Tk()
    app = LibraryManagementSystem(root)
    root.mainloop()
